"""
01_agreement_analysis.py
========================
Agreement analysis: LLM role persona vs. multidisciplinary tumour board.

This module evaluates the treatment-recommendation accuracy of three specialist
role personas (Surgeon, Oncologist, Radio-Oncologist) and a majority-vote
aggregation scheme against the reference standard defined by the institutional
multidisciplinary tumour board (MTB) record (``tumorboard_treatment``).

Outputs (saved to ``output/agreement_analysis`` and ``output/agreement_analysis/img``):
  - Overall and stratified agreement rates (Excel)
  - Bar charts per condition (PNG, 300 dpi)
  - Comparison heatmap (PNG)
  - Treatment summary table by diagnosis (Excel)
  - Word document with formatted analysis tables

Statistical tests performed:
  - Cochran's Q (omnibus)
  - Pairwise McNemar with Holm–Bonferroni correction
  - Wilson 95 % confidence intervals
"""

import os
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from scipy.stats import entropy

from config import (
    BAR_COLORS,
    ROLE_COLORS,
    DATA_FILE,
    OUTPUT_DIR_IMG_ROLE,
    OUTPUT_DIR_ROLE,
    RENAME_DICT,
    SHOW_PLOTS,
    TITLE_COLUMN_RENAME,
    VALUE_RENAME,
    COLUMNS_ANSWER,
    ROLE_PREFIX_MAP,
)

from utils import (
    calculate_correct_counts,
    calculate_correct_percentages,
    cochran_and_mcnemar,
    wilson_ci,
    parse_treatment_list_column
)

# ---------------------------------------------------------------------------
# Directory setup
# ---------------------------------------------------------------------------

TABLE_DIR = os.path.join(OUTPUT_DIR_ROLE, "agreement_analysis")
IMG_DIR = os.path.join(OUTPUT_DIR_ROLE, "agreement_analysis/img")
os.makedirs(TABLE_DIR, exist_ok=True)
os.makedirs(IMG_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------
df = pd.read_csv(DATA_FILE)

# ----------------------------------------------------------
# Ensure tumorboard_treatment is parsed as list
# and derive primary treatment (first element)
# ----------------------------------------------------------
df = parse_treatment_list_column(
    df,
    column="tumorboard_treatment",
    primary_output_col="tumorboard_primary_treatment"
)

# Columns to compare against the reference (excluding the reference itself)
comparison_cols = COLUMNS_ANSWER[1:]

# ---------------------------------------------------------------------------
# Overall agreement rates
# ---------------------------------------------------------------------------
percentages = calculate_correct_percentages(df, comparison_cols, RENAME_DICT)
counts      = calculate_correct_counts(df, comparison_cols)
ci_table    = wilson_ci(df, comparison_cols, RENAME_DICT)

print("\n=== Overall Agreement Rates (95 % Wilson CI) ===")
print(ci_table.to_string(index=False))

ci_table.to_excel(f"{TABLE_DIR}/overall_agreement_ci.xlsx", index=False)

# ---------------------------------------------------------------------------
# Statistical tests
# ---------------------------------------------------------------------------
comp_binary_cols = [f"{c}_treatment_concordance" for c in comparison_cols]

results = cochran_and_mcnemar(df, comp_binary_cols)

print("=== Cochran's Q Test ===")
q = results["cochran"]

print(
    f"Q = {q['Q']:.3f}, "
    f"df = {q['df']}, "
    f"p = {q['pvalue']:.4f}, "
    f"reject = {q['reject']}"
)

print("\n=== Pairwise McNemar Matrix ===")
print(results["pairwise_matrix"])

results["pairwise_matrix"].to_excel(
    f"{TABLE_DIR}/mcnemar_matrix.xlsx"
)


# ---------------------------------------------------------------------------
# Per-tumour Wilson CIs
# ---------------------------------------------------------------------------
print("\n=== 95 % CI per Tumour Type (Wilson method) ===")
ci_by_tumor_rows = []
for tumor in df["tumour_type"].unique():
    sub = df[df["tumour_type"] == tumor]
    n   = len(sub)
    label = VALUE_RENAME.get(tumor, tumor)
    for col in comparison_cols:
        correct = int(sub[f"{col}_treatment_concordance"].sum())
        p = correct / n if n > 0 else 0.0
        from statsmodels.stats.proportion import proportion_confint
        ci_low, ci_high = proportion_confint(correct, n, alpha=0.05, method="wilson")
        ci_by_tumor_rows.append(dict(
            tumour=label, model=RENAME_DICT.get(col, col),
            n=n, correct=correct,
            proportion=round(p, 4),
            ci_low=round(ci_low, 4), ci_high=round(ci_high, 4),
        ))

ci_by_tumor_df = pd.DataFrame(ci_by_tumor_rows)
ci_by_tumor_df.to_excel(f"{TABLE_DIR}/ci_by_tumor_type.xlsx", index=False)
print(ci_by_tumor_df.to_string(index=False))


# ===========================================================================
# Visualisation helpers
# ===========================================================================

def _save_or_show(path: str) -> None:
    """Save figure to *path* at 300 dpi; optionally display interactively."""
    plt.savefig(path, dpi=300, bbox_inches="tight")
    print(f"Figure saved → {path}")
    if SHOW_PLOTS:
        plt.show()
    plt.close()


def plot_concordance_heatmap(df: pd.DataFrame, comparison_cols: list,
                            save_dir: str) -> None:
    """
    Binary heatmap: green = correct, red = incorrect, per patient × model.

    Parameters
    ----------
    df :
        DataFrame with ``{col}_treatment_concordance`` columns.
    comparison_cols :
        Model base-name columns.
    save_dir :
        Directory to write the PNG file.
    """
    matrix = pd.DataFrame(
        {col: df[f"{col}_treatment_concordance"].values for col in comparison_cols}
    )
    matrix.index = range(1, len(matrix) + 1)

    plt.figure(figsize=(12, 8))
    sns.heatmap(
        matrix,
        cmap=sns.color_palette("RdYlGn", as_cmap=True),
        linewidths=0.5,
        linecolor="white",
        cbar=False,
    )
    plt.xlabel("Model", fontsize=11)
    plt.ylabel("Case Index", fontsize=11)
    plt.title("Agreement Heatmap: Correct (green) vs Incorrect (red)", fontsize=13)
    plt.xticks(
        ticks=np.arange(len(comparison_cols)) + 0.5,
        labels=[RENAME_DICT.get(c, c) for c in comparison_cols],
        rotation=40, ha="right",
    )
    _save_or_show(f"{save_dir}/agreement_heatmap.png")


def plot_overall_bar(percentages: dict, rename_dict: dict,
                     title: str, save_dir: str) -> None:
    """
    Grouped bar chart of overall agreement rates with percentage annotations.

    Parameters
    ----------
    percentages :
        Dict of ``{display_label: float}``.
    rename_dict :
        Label mapping (used for x-axis tick labels).
    title :
        Plot title; also used to construct the filename.
    save_dir :
        Directory to write the PNG file.
    """
    labels  = list(percentages.keys())
    values  = list(percentages.values())
    colors  = [BAR_COLORS[i % len(BAR_COLORS)] for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(labels, values, color=colors, edgecolor="white")

    for bar, val in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 1.5,
            f"{val:.1f}%",
            ha="center", va="bottom", fontsize=10,
        )

    ax.set_xlabel("Model / Framework", fontsize=11)
    ax.set_ylabel("Agreement Rate (%)", fontsize=11)
    ax.set_title(title, fontsize=13)
    ax.set_ylim(0, 115)
    plt.xticks(rotation=35, ha="right")
    plt.tight_layout()

    safe_title = title.replace("/", "-")
    _save_or_show(f"{save_dir}/role_{safe_title}.png")


def plot_sub_analysis(df: pd.DataFrame, column_name: str,
                      comparison_cols: list, rename_dict: dict,
                      save_dir: str) -> None:
    """
    Produce one agreement-rate bar chart per unique category in *column_name*.

    Parameters
    ----------
    df :
        DataFrame with comparison columns already computed.
    column_name :
        Stratification variable (e.g. ``"tumour_type"``, ``"presentation"``).
    comparison_cols :
        Model base-name columns.
    rename_dict :
        Label mapping.
    save_dir :
        Output directory.
    """
    for value in df[column_name].unique():
        sub = df[df[column_name] == value]
        if sub.empty:
            continue
        perc = calculate_correct_percentages(sub, comparison_cols, rename_dict)
        label = VALUE_RENAME.get(value, value)
        title = (f"Agreement Rate — "
                 f"{TITLE_COLUMN_RENAME.get(column_name, column_name)} = "
                 f"{label} (N = {len(sub)})")
        plot_overall_bar(perc, rename_dict, title, save_dir)

# ---------------------------------------------------------------------------
# Run visualisations
# ---------------------------------------------------------------------------
plot_concordance_heatmap(df, comparison_cols, IMG_DIR)
plot_overall_bar(percentages, RENAME_DICT, "Overall Agreement Rates", IMG_DIR)
plot_sub_analysis(df, "tumour_type",           comparison_cols, RENAME_DICT, IMG_DIR)
plot_sub_analysis(df, "presentation",                     comparison_cols, RENAME_DICT, IMG_DIR)
plot_sub_analysis(df, "tumorboard_primary_treatment", comparison_cols, RENAME_DICT, IMG_DIR)


# ===========================================================================
# Treatment distribution visualisation
# ===========================================================================

def create_stacked_bar_treatment(df: pd.DataFrame,
                                  diagnosis_col: str, evwv_col: str,
                                  treatment_col: str, save_dir: str) -> None:
    """
    Stacked bar chart: treatment recommendation counts by diagnosis × consultation type.

    Saves the underlying frequency table to Excel for supplementary data.
    """
    grouped = (
        df.groupby([diagnosis_col, evwv_col, treatment_col])
        .size()
        .unstack(fill_value=0)
        .reset_index()
    )
    grouped[diagnosis_col] = grouped[diagnosis_col].replace(VALUE_RENAME)
    grouped[evwv_col]      = grouped[evwv_col].replace(VALUE_RENAME)
    grouped.to_excel(f"{TABLE_DIR}/recommendation_per_tumor_type.xlsx", index=False)

    grouped["combined_label"] = (
        grouped[diagnosis_col].astype(str) + " — " + grouped[evwv_col].astype(str)
    )
    grouped = grouped.set_index("combined_label")
    numeric_cols = [c for c in grouped.columns
                    if c not in [diagnosis_col, evwv_col, "combined_label"]]
    grouped = grouped[numeric_cols].apply(pd.to_numeric, errors="coerce").fillna(0)

    colors = sns.color_palette("Set2", len(numeric_cols))
    fig, ax = plt.subplots(figsize=(12, 6))
    bottom = np.zeros(len(grouped))
    for i, treat in enumerate(numeric_cols):
        vals = grouped[treat].values.astype(float)
        ax.bar(grouped.index, vals, label=treat, bottom=bottom, color=colors[i])
        bottom += vals

    ax.set_xlabel("Diagnosis — Consultation Type", fontsize=11)
    ax.set_ylabel("Count", fontsize=11)
    ax.set_title("Tumour Board Recommendations by Diagnosis and Consultation Type", fontsize=13)
    plt.xticks(rotation=55, ha="right")
    plt.legend(title="Recommendation", bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()
    _save_or_show(f"{save_dir}/stacked_treatment_distribution.png")


def create_treatment_summary_table(df: pd.DataFrame, diagnosis_col: str,
                                    evwv_col: str, treatment_col: str,
                                    save_path: str) -> pd.DataFrame:
    """
    Build a publication-ready treatment summary table: n (%) per diagnosis.

    Rows represent treatment categories (overall, First Presentation, FUP).
    Columns represent tumour type plus an "Overall" summary column.

    Parameters
    ----------
    df :
        Full patient DataFrame.
    diagnosis_col :
        Column name for tumour type.
    evwv_col :
        Column name for consultation type (EV / WV).
    treatment_col :
        Column with the treatment category.
    save_path :
        Full file path for the Excel output.

    Returns
    -------
    pandas.DataFrame
        The formatted summary table.
    """
    diagnoses  = sorted(df[diagnosis_col].dropna().unique())
    treatments = sorted(df[treatment_col].dropna().unique())

    row_index  = []
    for t in treatments:
        row_index += [f"{t}, n (%)", f"  - First Presentation", f"  - FUP"]

    col_order  = ["Overall"] + diagnoses
    table      = pd.DataFrame(index=row_index, columns=col_order)

    for col in col_order:
        subset = df if col == "Overall" else df[df[diagnosis_col] == col]
        total  = len(subset)
        fp_sub = subset[subset[evwv_col] == "EV"]
        fu_sub = subset[subset[evwv_col] == "WV"]

        for t in treatments:
            def fmt(sub, denom):
                c = (sub[treatment_col] == t).sum()
                p = c / denom * 100 if denom else 0.0
                return f"{c} ({p:.1f}%)"
            table.at[f"{t}, n (%)",          col] = fmt(subset, total)
            table.at[f"  - First Presentation", col] = fmt(fp_sub, len(fp_sub))
            table.at[f"  - FUP",             col] = fmt(fu_sub, len(fu_sub))

    # Prepend sample-size header row
    n_row = pd.DataFrame(
        {col: f"N = {len(df[df[diagnosis_col] == col])}" if col != "Overall"
              else f"N = {len(df)}"
         for col in col_order},
        index=[""],
    )
    table = pd.concat([n_row, table])
    table.to_excel(save_path)
    print(f"Treatment summary saved → {save_path}")
    return table


create_stacked_bar_treatment(df, "tumour_type", "presentation",
                              "tumorboard_primary_treatment", IMG_DIR)
create_treatment_summary_table(
    df,
    diagnosis_col="tumour_type",
    evwv_col="presentation",
    treatment_col="tumorboard_primary_treatment",
    save_path=f"{TABLE_DIR}/treatment_summary_by_diagnosis.xlsx",
)


# ===========================================================================
# Role-specific treatment distribution
# ===========================================================================

def build_role_treatment_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    records = []

    for role, prefix in ROLE_PREFIX_MAP.items():
        col = f"{prefix}_treatment"

        if col not in df.columns:
            continue

        for val in df[col].dropna():
            records.append({
                "Role": role,
                "Treatment": VALUE_RENAME.get(val, val)
            })

    return pd.DataFrame(records)


def plot_role_treatment_distribution(stats: pd.DataFrame, save_dir: str) -> None:
    """
    Grouped bar chart: treatment recommendation percentages by specialist role.

    Parameters
    ----------
    stats :
        Long-format DataFrame with columns: Role, Treatment, Count, Percentage.
    save_dir :
        Output directory.
    """
    role_colors = ROLE_COLORS
    plt.figure(figsize=(14, 6))
    sns.barplot(
        data=stats, x="Treatment", y="Percentage",
        hue="Role", dodge=True, palette=role_colors,
    )
    plt.ylabel("Percentage of Recommendations (%)", fontsize=11)
    plt.xlabel("Treatment Category", fontsize=11)
    plt.title("Treatment Recommendations by Specialist Role", fontsize=13)
    plt.xticks(rotation=40, ha="right")
    plt.ylim(0, 100)
    plt.legend(title="Role", bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()
    _save_or_show(f"{save_dir}/role_treatment_distribution.png")


role_treatment_df    = build_role_treatment_dataframe(df)
role_treatment_stats = (
    role_treatment_df
    .groupby(["Role", "Treatment"])
    .size()
    .reset_index(name="Count")
)
role_treatment_stats["Percentage"] = (
    role_treatment_stats
    .groupby("Role")["Count"]
    .transform(lambda x: x / x.sum() * 100)
)
plot_role_treatment_distribution(role_treatment_stats, IMG_DIR)

print('=== Role Treatment Statistics ===')
print(role_treatment_stats)
role_treatment_stats.to_excel(f"{TABLE_DIR}/role_treatment_percentages.xlsx", index=False)


role_treatment_cols = [
    "F3_persona_surgeon_treatment",
    "F4_persona_medical_oncologist_treatment",
    "F5_persona_radiation_oncologist_treatment",
]

def shannon_entropy_row(row):
    vals = [row[c] for c in role_treatment_cols if pd.notna(row[c])]
    if not vals: return np.nan
    counts = pd.Series(vals).value_counts(normalize=True)
    if len(counts) == 0:
        return np.nan
    return float(entropy(counts, base=2))

df['treatment_entropy_bits'] = df.apply(shannon_entropy_row, axis=1)
print(df['treatment_entropy_bits'].describe())
print(df.groupby('tumour_type')['treatment_entropy_bits'].mean())



# ===========================================================================
# Word document tables
# ===========================================================================

def _add_table_to_doc(doc: Document, df: pd.DataFrame,
                       column_name: str, comparison_cols: list,
                       rename_dict: dict) -> None:
    """Append a stratified agreement-rate table to *doc*."""
    unique_vals   = df[column_name].unique()
    title_label   = TITLE_COLUMN_RENAME.get(column_name, column_name)
    renamed_cols  = [rename_dict.get(c, c) for c in comparison_cols]

    doc.add_paragraph(f"Table: Agreement Rate by {title_label}")

    table = doc.add_table(rows=1, cols=len(renamed_cols) + 1)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = title_label
    for i, lbl in enumerate(renamed_cols):
        hdr[i + 1].text = lbl

    for value in unique_vals:
        sub     = df[df[column_name] == value]
        n       = len(sub)
        label   = VALUE_RENAME.get(value, value)
        cnts    = calculate_correct_counts(sub, comparison_cols)
        percs   = calculate_correct_percentages(sub, comparison_cols, rename_dict)

        row_cells = table.add_row().cells
        row_cells[0].text = f"{label}  (N = {n})"
        for j, col_lbl in enumerate(renamed_cols):
            cnt = cnts.get(list(rename_dict.keys())[j] if rename_dict else col_lbl, 0)
            prc = percs.get(col_lbl, 0.0)
            row_cells[j + 1].text = f"{int(cnt)} ({prc:.1f}%)"

    doc.add_paragraph("")


doc = Document()
_add_table_to_doc(doc, df, "tumour_type",           comparison_cols, RENAME_DICT)
_add_table_to_doc(doc, df, "presentation",                     comparison_cols, RENAME_DICT)
_add_table_to_doc(doc, df, "tumorboard_primary_treatment", comparison_cols, RENAME_DICT)
doc.save(f"{TABLE_DIR}/Analysis_Tables_role.docx")
print(f"Word tables saved → {TABLE_DIR}/Analysis_Tables_role.docx")

print("\n=== Agreement Analysis Complete ===")
